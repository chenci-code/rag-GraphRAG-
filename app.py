from __future__ import annotations

import asyncio
import gc
import io
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import uuid
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

import pandas as pd
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

import graphrag.api as graphrag_api
from graphrag.config.embeddings import text_unit_text_embedding
from graphrag.api.index import build_index
from graphrag.config.enums import IndexingMethod
from graphrag.config.load_config import load_config
from graphrag.config.models.graph_rag_config import GraphRagConfig
from graphrag.data_model.data_reader import DataReader
from graphrag.index.validate_config import validate_config_names
from graphrag.prompts.query.basic_search_system_prompt import BASIC_SEARCH_SYSTEM_PROMPT
from graphrag.utils.api import get_embedding_store, load_search_prompt
from graphrag_llm.completion.completion_factory import create_completion
from graphrag_llm.embedding.embedding_factory import create_embedding
from graphrag_llm.utils import CompletionMessagesBuilder
from graphrag_storage import create_storage
from graphrag_storage.tables.table_provider_factory import create_table_provider

from mysql_metadata import MySQLPermissionStore, sync_graphrag_metadata_to_mysql
from neo4j_cleanup import Neo4jFileCleaner

QueryMethod = Literal["basic", "local", "global", "drift", "hybrid"]
RebuildMode = Literal["full"]
ProcessingStatus = Literal["PROCESSING", "SUCCESS", "FAILED"]


class QueryRequest(BaseModel):
    question: str = Field(min_length=1, description="The user question to ask.")
    method: QueryMethod = Field(
        default="local", description="The GraphRAG search method to use."
    )
    community_level: int | None = Field(
        default=2,
        description="Community level for local/global/drift methods.",
    )
    dynamic_community_selection: bool = Field(
        default=False,
        description="Enable dynamic community selection for global search.",
    )
    response_type: str = Field(
        default="客服短答：1段话或3-5个要点，去掉来源引用",
        description="Free-form response format hint passed to GraphRAG.",
    )
    include_context: bool = Field(
        default=True,
        description="Whether to include the raw GraphRAG context payload.",
    )
    verbose: bool = Field(
        default=False,
        description="Enable GraphRAG verbose query logging.",
    )
    hybrid_top_k: int = Field(
        default=10,
        ge=1,
        le=50,
        description="Number of fused chunks to include for hybrid search.",
    )
    hybrid_vector_k: int = Field(
        default=20,
        ge=1,
        le=100,
        description="Number of vector candidates to retrieve for hybrid search.",
    )
    hybrid_keyword_k: int = Field(
        default=20,
        ge=1,
        le=100,
        description="Number of keyword/BM25 candidates to retrieve for hybrid search.",
    )
    hybrid_vector_weight: float = Field(
        default=0.6,
        ge=0,
        le=1,
        description="RRF weight for vector candidates in hybrid search.",
    )
    hybrid_keyword_weight: float = Field(
        default=0.4,
        ge=0,
        le=1,
        description="RRF weight for keyword/BM25 candidates in hybrid search.",
    )
    hybrid_graph_k: int = Field(
        default=20,
        ge=1,
        le=100,
        description="Number of graph-neighborhood candidates to retrieve for hybrid search.",
    )
    hybrid_graph_weight: float = Field(
        default=0.5,
        ge=0,
        le=1,
        description="RRF weight for graph-neighborhood candidates in hybrid search.",
    )
    tenant_id: str | None = Field(
        default=None,
        description="Tenant scope used by metadata/permission stores.",
    )
    conversation_id: str | None = Field(
        default=None,
        description="Conversation scope used to isolate uploaded conversation files.",
    )
    enforce_permissions: bool = Field(
        default=False,
        description="Filter retrieved text units through the configured permission store.",
    )
    user_id: str | None = Field(
        default=None,
        description="Current user ID for permission filtering.",
    )
    department_id: str | None = Field(
        default=None,
        description="Current department ID for permission filtering.",
    )
    role_ids: list[str] = Field(
        default_factory=list,
        description="Current role IDs for permission filtering.",
    )


class QueryResponse(BaseModel):
    method: QueryMethod
    answer: Any
    context: Any | None = None


class FileInfo(BaseModel):
    filename: str
    size: int
    updated_at: float
    tenant_id: str = "default"
    conversation_id: str | None = None
    department_id: str = "default"
    visibility: str = "department"
    file_id: str | None = None
    status: ProcessingStatus = "SUCCESS"
    error: str | None = None


class FileWriteRequest(BaseModel):
    filename: str = Field(min_length=1, description="Target filename under input/.")
    content: str = Field(description="UTF-8 text content to store in the file.")
    overwrite: bool = Field(
        default=False,
        description="Whether to replace an existing file with the same name.",
    )
    tenant_id: str = Field(
        default="default",
        description="Tenant that owns this file.",
    )
    conversation_id: str | None = Field(
        default=None,
        description="Optional conversation scope that owns this file.",
    )
    department_id: str = Field(
        default="default",
        description="Department metadata used by permission sync.",
    )
    visibility: str = Field(
        default="department",
        description="Visibility metadata used by permission sync.",
    )


class FileMutationResponse(BaseModel):
    filename: str
    action: Literal["created", "replaced", "deleted"]
    rebuild_mode: RebuildMode
    stats: dict[str, Any]
    source_filename: str | None = None
    converted: bool = False
    conversion_type: str | None = None
    warnings: list[str] = Field(default_factory=list)
    tenant_id: str = "default"
    conversation_id: str | None = None
    file_id: str | None = None
    status: ProcessingStatus = "SUCCESS"
    task_id: str | None = None
    error: str | None = None


class FilePreviewResponse(BaseModel):
    filename: str
    source_filename: str
    converted: bool
    conversion_type: str
    warnings: list[str] = Field(default_factory=list)
    content: str


class RebuildRequest(BaseModel):
    verbose: bool = Field(
        default=False,
        description="Enable verbose indexing logs during rebuild.",
    )
    sync_mysql: bool = Field(
        default=True,
        description="Sync GraphRAG metadata into MySQL after a successful rebuild.",
    )


class OcrConfigResponse(BaseModel):
    provider: str
    enabled: bool
    model: str
    api_key_configured: bool
    api_key_source: str
    timeout_seconds: int
    command_configured: bool
    script_exists: bool
    message: str


class ChatStreamRequest(BaseModel):
    conversation_id: str = Field(min_length=1)
    message: str = Field(min_length=1)
    tenant_id: str = Field(default="default")
    method: QueryMethod = "hybrid"
    include_context: bool = True
    response_type: str = "客服短答：1段话或3-5个要点，去掉来源引用"
    hybrid_top_k: int = Field(default=10, ge=1, le=50)
    hybrid_graph_k: int = Field(default=20, ge=1, le=100)


@dataclass
class LoadedTables:
    documents: pd.DataFrame
    entities: pd.DataFrame
    communities: pd.DataFrame
    community_reports: pd.DataFrame
    text_units: pd.DataFrame
    relationships: pd.DataFrame
    covariates: pd.DataFrame | None


@dataclass
class NormalizedUpload:
    filename: str
    content: str
    converted: bool
    conversion_type: str
    warnings: list[str]


@dataclass
class DetectedFileType:
    kind: Literal["text", "markdown", "csv", "pdf", "docx", "xlsx"]
    extension: str
    detected_by: str


class _MemoryUploadFile:
    def __init__(self, filename: str, payload: bytes):
        self.filename = filename
        self._payload = payload

    async def read(self) -> bytes:
        return self._payload


class GraphRAGService:
    def __init__(self, root: Path):
        self.root = root.resolve()
        load_dotenv(self.root / ".env", override=False)
        self.input_dir = self.root / "input"
        self.output_dir = self.root / "output"
        self.output_backup_dir = self.root / "output.__backup__"
        self.metadata_overrides_path = self.root / "metadata_overrides.json"
        self.processing_tasks_path = self.root / "processing_tasks.json"
        self.config: GraphRagConfig | None = None
        self.tables: LoadedTables | None = None
        self._lock = asyncio.Lock()
        self.permission_store = MySQLPermissionStore.from_env()
        self.neo4j_cleaner = Neo4jFileCleaner.from_env()

    async def reload(self) -> dict[str, Any]:
        async with self._lock:
            return await self._reload_unlocked()

    async def query(self, request: QueryRequest) -> QueryResponse:
        async with self._lock:
            if self.config is None or self.tables is None:
                await self._reload_unlocked()

            assert self.config is not None
            assert self.tables is not None

            if request.method == "basic":
                text_units = self._scoped_text_units(request)
                answer, context = await graphrag_api.basic_search(
                    config=self.config,
                    text_units=text_units,
                    response_type=request.response_type,
                    query=request.question,
                    verbose=request.verbose,
                )
            elif request.method == "local":
                self._reject_unscoped_graph_method(request)
                community_level = (
                    request.community_level if request.community_level is not None else 2
                )
                answer, context = await graphrag_api.local_search(
                    config=self.config,
                    entities=self.tables.entities,
                    communities=self.tables.communities,
                    community_reports=self.tables.community_reports,
                    text_units=self.tables.text_units,
                    relationships=self.tables.relationships,
                    covariates=self.tables.covariates,
                    community_level=community_level,
                    response_type=request.response_type,
                    query=request.question,
                    verbose=request.verbose,
                )
            elif request.method == "global":
                self._reject_unscoped_graph_method(request)
                answer, context = await graphrag_api.global_search(
                    config=self.config,
                    entities=self.tables.entities,
                    communities=self.tables.communities,
                    community_reports=self.tables.community_reports,
                    community_level=request.community_level,
                    dynamic_community_selection=request.dynamic_community_selection,
                    response_type=request.response_type,
                    query=request.question,
                    verbose=request.verbose,
                )
            elif request.method == "hybrid":
                answer, context = await self._hybrid_search_unlocked(request)
            else:
                self._reject_unscoped_graph_method(request)
                community_level = (
                    request.community_level if request.community_level is not None else 2
                )
                answer, context = await graphrag_api.drift_search(
                    config=self.config,
                    entities=self.tables.entities,
                    communities=self.tables.communities,
                    community_reports=self.tables.community_reports,
                    text_units=self.tables.text_units,
                    relationships=self.tables.relationships,
                    community_level=community_level,
                    response_type=request.response_type,
                    query=request.question,
                    verbose=request.verbose,
                )

            return QueryResponse(
                method=request.method,
                answer=clean_customer_answer(to_jsonable(answer)),
                context=to_jsonable(context) if request.include_context else None,
            )

    async def _hybrid_search_unlocked(
        self, request: QueryRequest
    ) -> tuple[str, dict[str, pd.DataFrame]]:
        context_text, context_payload = self._hybrid_context_unlocked(request)
        if not context_text:
            return "", context_payload
        answer = await self._generate_basic_answer(
            query=request.question.strip(),
            context_text=context_text,
            response_type=request.response_type,
        )
        return answer, context_payload

    def _hybrid_context_unlocked(
        self, request: QueryRequest
    ) -> tuple[str, dict[str, pd.DataFrame]]:
        assert self.config is not None
        assert self.tables is not None

        query = request.question.strip()
        text_units = self._scoped_text_units(request)
        if query == "" or text_units.empty:
            return "", {"sources": pd.DataFrame(columns=["id", "text"])}

        allowed_text_unit_ids = self._allowed_text_unit_ids(request)
        if allowed_text_unit_ids is not None:
            text_units = text_units[
                text_units["id"].astype(str).isin(allowed_text_unit_ids)
            ].copy()
            if text_units.empty:
                return "", {"sources": pd.DataFrame(columns=["id", "text"])}

        id_to_row = {
            str(row["id"]): row for row in text_units.to_dict(orient="records")
        }
        vector_ranked_ids = self._vector_rank_text_units(
            query=query,
            k=request.hybrid_vector_k,
            allowed_ids=allowed_text_unit_ids,
        )
        keyword_ranked_ids = self._keyword_rank_text_units(
            text_units=text_units,
            query=query,
            k=request.hybrid_keyword_k,
        )
        graph_ranked_ids, graph_rows = self._graph_rank_text_units(
            text_units=text_units,
            query=query,
            k=request.hybrid_graph_k,
            allowed_ids=allowed_text_unit_ids,
        )
        fused_ids = self._rrf_fuse(
            ranked_lists=[
                (vector_ranked_ids, request.hybrid_vector_weight),
                (keyword_ranked_ids, request.hybrid_keyword_weight),
                (graph_ranked_ids, request.hybrid_graph_weight),
            ],
            limit=request.hybrid_top_k,
        )

        source_rows: list[dict[str, Any]] = []
        for rank, text_unit_id in enumerate(fused_ids, start=1):
            row = id_to_row.get(text_unit_id)
            if row is None:
                continue
            source_rows.append({
                "id": str(row.get("human_readable_id", rank)),
                "text": row.get("text", ""),
                "text_unit_id": text_unit_id,
                "hybrid_rank": rank,
                "vector_rank": self._rank_or_none(vector_ranked_ids, text_unit_id),
                "keyword_rank": self._rank_or_none(keyword_ranked_ids, text_unit_id),
                "graph_rank": self._rank_or_none(graph_ranked_ids, text_unit_id),
            })

        sources_df = pd.DataFrame(source_rows)
        if sources_df.empty:
            return "", {"sources": sources_df}

        graph_df = pd.DataFrame(graph_rows)
        context_rows = sources_df[["id", "text"]].to_dict(orient="records")
        if not graph_df.empty:
            for rank, row in enumerate(graph_df.to_dict(orient="records"), start=1):
                context_rows.append({
                    "id": f"graph:{rank}",
                    "text": row["text"],
                })

        context_df = pd.DataFrame(context_rows)
        context_text = context_df.to_csv(index=False, escapechar="\\", sep="|")
        return context_text, {"sources": sources_df, "graph_context": graph_df}

    def _allowed_text_unit_ids(self, request: QueryRequest) -> set[str] | None:
        local_ids = self._local_scope_text_unit_ids(request)
        if not request.enforce_permissions:
            return local_ids

        if not self.permission_store.enabled:
            msg = (
                "Permission filtering requested, but MySQL permission mode is not "
                "configured. Use tenant/conversation filters for local isolation, or "
                "set GRAPHRAG_PERMISSION_MODE=mysql and MySQL settings."
            )
            raise RuntimeError(msg)

        mysql_ids = self.permission_store.allowed_text_unit_ids(
            tenant_id=request.tenant_id,
            user_id=request.user_id,
            department_id=request.department_id,
            role_ids=request.role_ids,
        )
        if local_ids is None:
            return mysql_ids
        return local_ids & mysql_ids

    def _scoped_text_units(self, request: QueryRequest) -> pd.DataFrame:
        assert self.tables is not None
        text_units = self.tables.text_units.copy()
        scoped_ids = self._local_scope_text_unit_ids(request)
        if scoped_ids is None:
            return text_units
        return text_units[text_units["id"].astype(str).isin(scoped_ids)].copy()

    def _local_scope_text_unit_ids(self, request: QueryRequest) -> set[str] | None:
        assert self.tables is not None
        if not request.tenant_id and not request.conversation_id:
            return None

        document_scope = self._document_scope_map()
        if not document_scope:
            if request.tenant_id and request.tenant_id != "default":
                return set()
            if request.conversation_id:
                return set()
            return None

        allowed_document_ids: set[str] = set()
        for document_id, metadata in document_scope.items():
            if request.tenant_id and metadata.get("tenant_id", "default") != request.tenant_id:
                continue
            if (
                request.conversation_id
                and metadata.get("conversation_id") != request.conversation_id
            ):
                continue
            allowed_document_ids.add(document_id)

        if not allowed_document_ids:
            return set()

        return {
            str(row["id"])
            for row in self.tables.text_units.to_dict(orient="records")
            if str(row.get("document_id")) in allowed_document_ids
        }

    def _reject_unscoped_graph_method(self, request: QueryRequest) -> None:
        if request.tenant_id or request.conversation_id:
            msg = (
                f"{request.method} search cannot be safely tenant/conversation scoped "
                "in this local GraphRAG table layout. Use hybrid or basic for isolated "
                "customer-service conversations."
            )
            raise RuntimeError(msg)

    def _vector_rank_text_units(
        self,
        query: str,
        k: int,
        allowed_ids: set[str] | None = None,
    ) -> list[str]:
        assert self.config is not None
        embedding_settings = self.config.get_embedding_model_config(
            self.config.basic_search.embedding_model_id
        )
        embedding_model = create_embedding(embedding_settings)
        embedding_store = get_embedding_store(
            config=self.config.vector_store,
            embedding_name=text_unit_text_embedding,
        )
        vector_k = k if allowed_ids is None else max(k * 5, k)
        results = embedding_store.similarity_search_by_text(
            text=query,
            text_embedder=lambda text: (
                embedding_model.embedding(input=[text]).first_embedding
            ),
            k=vector_k,
            include_vectors=False,
        )
        ranked_ids = [str(result.document.id) for result in results]
        if allowed_ids is not None:
            ranked_ids = [doc_id for doc_id in ranked_ids if doc_id in allowed_ids]
        return ranked_ids[:k]

    def _keyword_rank_text_units(
        self,
        text_units: pd.DataFrame,
        query: str,
        k: int,
    ) -> list[str]:
        docs = [
            (str(row["id"]), _tokenize_for_keyword_search(str(row.get("text", ""))))
            for row in text_units.to_dict(orient="records")
        ]
        query_terms = _tokenize_for_keyword_search(query)
        if not query_terms:
            return []

        doc_freq: defaultdict[str, int] = defaultdict(int)
        doc_lengths: dict[str, int] = {}
        term_counts: dict[str, Counter[str]] = {}
        for doc_id, tokens in docs:
            counts = Counter(tokens)
            term_counts[doc_id] = counts
            doc_lengths[doc_id] = len(tokens)
            for term in counts:
                doc_freq[term] += 1

        total_docs = len(docs)
        avg_doc_len = sum(doc_lengths.values()) / total_docs if total_docs else 0
        if avg_doc_len == 0:
            return []

        k1 = 1.5
        b = 0.75
        scores: dict[str, float] = defaultdict(float)
        for term in query_terms:
            df = doc_freq.get(term, 0)
            if df == 0:
                continue
            idf = math.log(1 + (total_docs - df + 0.5) / (df + 0.5))
            for doc_id, counts in term_counts.items():
                freq = counts.get(term, 0)
                if freq == 0:
                    continue
                denom = freq + k1 * (1 - b + b * doc_lengths[doc_id] / avg_doc_len)
                scores[doc_id] += idf * freq * (k1 + 1) / denom

        return [
            doc_id
            for doc_id, _ in sorted(scores.items(), key=lambda item: item[1], reverse=True)[
                :k
            ]
        ]

    def _graph_rank_text_units(
        self,
        text_units: pd.DataFrame,
        query: str,
        k: int,
        allowed_ids: set[str] | None = None,
        max_hops: int = 2,
    ) -> tuple[list[str], list[dict[str, Any]]]:
        assert self.tables is not None
        if self.tables.entities.empty or self.tables.relationships.empty:
            return [], []

        available_ids = {str(row["id"]) for row in text_units.to_dict(orient="records")}
        if allowed_ids is not None:
            available_ids &= allowed_ids
        if not available_ids:
            return [], []

        entity_scores = self._score_graph_anchor_entities(query)
        if not entity_scores:
            return self._rank_relationships_by_keyword(
                query=query,
                relationships=self.tables.relationships,
                available_ids=available_ids,
                k=k,
            )

        adjacency: defaultdict[str, list[dict[str, Any]]] = defaultdict(list)
        for row in self.tables.relationships.to_dict(orient="records"):
            source = str(row.get("source", ""))
            target = str(row.get("target", ""))
            if not source or not target:
                continue
            adjacency[source].append(row)
            adjacency[target].append(row)

        relationship_scores: dict[str, tuple[float, dict[str, Any]]] = {}
        frontier = {title: score for title, score in entity_scores[:12]}
        visited_entities: set[str] = set()

        for hop in range(max_hops):
            next_frontier: dict[str, float] = {}
            for entity_title, anchor_score in frontier.items():
                if entity_title in visited_entities:
                    continue
                visited_entities.add(entity_title)
                for rel in adjacency.get(entity_title, []):
                    rel_id = str(rel.get("id", ""))
                    if not rel_id:
                        continue
                    source = str(rel.get("source", ""))
                    target = str(rel.get("target", ""))
                    other = target if source == entity_title else source
                    score = (
                        anchor_score * (0.65**hop)
                        + self._keyword_overlap_score(
                            query,
                            " ".join([
                                source,
                                target,
                                str(rel.get("description", "")),
                            ]),
                        )
                        + float(rel.get("weight", 0) or 0) * 0.03
                    )
                    existing = relationship_scores.get(rel_id)
                    if existing is None or score > existing[0]:
                        relationship_scores[rel_id] = (score, rel)
                    if other not in visited_entities:
                        next_frontier[other] = max(next_frontier.get(other, 0), score)
            frontier = next_frontier
            if not frontier:
                break

        ranked_relationships = [
            rel
            for _, rel in sorted(
                relationship_scores.values(),
                key=lambda item: item[0],
                reverse=True,
            )
        ]
        return self._relationship_rows_to_text_units(
            ranked_relationships=ranked_relationships,
            available_ids=available_ids,
            k=k,
        )

    def _score_graph_anchor_entities(self, query: str) -> list[tuple[str, float]]:
        assert self.tables is not None
        normalized_query = query.lower()
        scored: list[tuple[str, float]] = []
        for row in self.tables.entities.to_dict(orient="records"):
            title = str(row.get("title", "")).strip()
            if not title:
                continue
            title_lower = title.lower()
            score = 0.0
            if title_lower and title_lower in normalized_query:
                score += 8.0 + min(len(title), 12) * 0.2
            score += self._keyword_overlap_score(
                query,
                " ".join([
                    title,
                    str(row.get("type", "")),
                    str(row.get("description", "")),
                ]),
            )
            if score > 0:
                scored.append((title, score))
        return sorted(scored, key=lambda item: item[1], reverse=True)

    def _rank_relationships_by_keyword(
        self,
        query: str,
        relationships: pd.DataFrame,
        available_ids: set[str],
        k: int,
    ) -> tuple[list[str], list[dict[str, Any]]]:
        scored: list[tuple[float, dict[str, Any]]] = []
        for row in relationships.to_dict(orient="records"):
            score = self._keyword_overlap_score(
                query,
                " ".join([
                    str(row.get("source", "")),
                    str(row.get("target", "")),
                    str(row.get("description", "")),
                ]),
            )
            if score > 0:
                scored.append((score, row))
        ranked_relationships = [
            row for _, row in sorted(scored, key=lambda item: item[0], reverse=True)
        ]
        return self._relationship_rows_to_text_units(
            ranked_relationships=ranked_relationships,
            available_ids=available_ids,
            k=k,
        )

    def _relationship_rows_to_text_units(
        self,
        ranked_relationships: list[dict[str, Any]],
        available_ids: set[str],
        k: int,
    ) -> tuple[list[str], list[dict[str, Any]]]:
        text_unit_scores: defaultdict[str, float] = defaultdict(float)
        graph_rows: list[dict[str, Any]] = []

        for rank, rel in enumerate(ranked_relationships, start=1):
            rel_text_unit_ids = [
                str(item)
                for item in _coerce_list(rel.get("text_unit_ids"))
                if str(item) in available_ids
            ]
            if not rel_text_unit_ids:
                continue

            score = 1 / rank
            for text_unit_id in rel_text_unit_ids:
                text_unit_scores[text_unit_id] += score

            source = str(rel.get("source", ""))
            target = str(rel.get("target", ""))
            description = str(rel.get("description", ""))
            graph_rows.append({
                "relationship_id": str(rel.get("id", "")),
                "source": source,
                "target": target,
                "description": description,
                "text_unit_ids": rel_text_unit_ids,
                "text": f"{source} -> {target}: {description}",
            })
            if len(graph_rows) >= k:
                break

        ranked_ids = [
            text_unit_id
            for text_unit_id, _ in sorted(
                text_unit_scores.items(),
                key=lambda item: item[1],
                reverse=True,
            )[:k]
        ]
        return ranked_ids, graph_rows

    def _keyword_overlap_score(self, query: str, text: str) -> float:
        query_terms = Counter(_tokenize_for_keyword_search(query))
        if not query_terms:
            return 0.0
        text_terms = Counter(_tokenize_for_keyword_search(text))
        if not text_terms:
            return 0.0
        score = 0.0
        for term, query_count in query_terms.items():
            score += min(query_count, text_terms.get(term, 0))
        return score / max(sum(query_terms.values()), 1)

    def _rrf_fuse(
        self,
        ranked_lists: list[tuple[list[str], float]],
        limit: int,
        rank_constant: int = 60,
    ) -> list[str]:
        if all(weight <= 0 for _, weight in ranked_lists):
            ranked_lists = [(ranked_ids, 1.0) for ranked_ids, _ in ranked_lists]

        scores: defaultdict[str, float] = defaultdict(float)
        for ranked_ids, weight in ranked_lists:
            for rank, doc_id in enumerate(ranked_ids, start=1):
                scores[doc_id] += weight / (rank_constant + rank)
        return [
            doc_id
            for doc_id, _ in sorted(scores.items(), key=lambda item: item[1], reverse=True)[
                :limit
            ]
        ]

    @staticmethod
    def _rank_or_none(ranked_ids: list[str], doc_id: str) -> int | None:
        try:
            return ranked_ids.index(doc_id) + 1
        except ValueError:
            return None

    async def _generate_basic_answer(
        self,
        query: str,
        context_text: str,
        response_type: str,
    ) -> str:
        assert self.config is not None
        model_settings = self.config.get_completion_model_config(
            self.config.basic_search.completion_model_id
        )
        chat_model = create_completion(model_settings)
        prompt = load_search_prompt(self.config.basic_search.prompt)
        system_prompt = prompt or BASIC_SEARCH_SYSTEM_PROMPT
        search_prompt = system_prompt.format(
            context_data=context_text,
            response_type=response_type,
        )
        messages = (
            CompletionMessagesBuilder()
            .add_system_message(search_prompt)
            .add_user_message(query)
            .build()
        )
        response = await chat_model.completion_async(
            messages=messages,
            stream=False,
            **model_settings.call_args,
        )
        return clean_customer_answer(response.choices[0].message.content or "")

    async def _stream_basic_answer(
        self,
        query: str,
        context_text: str,
        response_type: str,
    ):
        assert self.config is not None
        model_settings = self.config.get_completion_model_config(
            self.config.basic_search.completion_model_id
        )
        chat_model = create_completion(model_settings)
        prompt = load_search_prompt(self.config.basic_search.prompt)
        system_prompt = prompt or BASIC_SEARCH_SYSTEM_PROMPT
        search_prompt = system_prompt.format(
            context_data=context_text,
            response_type=response_type,
        )
        messages = (
            CompletionMessagesBuilder()
            .add_system_message(search_prompt)
            .add_user_message(query)
            .build()
        )
        response_stream = await chat_model.completion_async(
            messages=messages,
            stream=True,
            **model_settings.call_args,
        )
        async for chunk in response_stream:
            yield chunk.choices[0].delta.content or ""

    async def list_files(self) -> list[FileInfo]:
        async with self._lock:
            self.input_dir.mkdir(parents=True, exist_ok=True)
            overrides = self._load_metadata_overrides()
            tasks = self._load_processing_tasks()
            files = [
                FileInfo(
                    filename=path.name,
                    size=path.stat().st_size,
                    updated_at=path.stat().st_mtime,
                    status=self._latest_file_task_status(path.name, tasks),
                    error=self._latest_file_task_error(path.name, tasks),
                    **self._file_metadata_for_response(path.name, overrides),
                )
                for path in sorted(self.input_dir.iterdir())
                if path.is_file()
            ]
            existing_names = {file.filename for file in files}
            for task in tasks.values():
                if task.get("status") != "PROCESSING":
                    continue
                filename = str(task.get("filename", ""))
                if not filename or filename in existing_names:
                    continue
                files.append(
                    FileInfo(
                        filename=filename,
                        size=0,
                        updated_at=float(task.get("updated_at", task.get("created_at", 0))),
                        tenant_id=str(task.get("tenant_id", "default")),
                        conversation_id=task.get("conversation_id"),
                        department_id=str(task.get("department_id", "default")),
                        visibility=str(task.get("visibility", "department")),
                        file_id=filename,
                        status=task.get("status", "PROCESSING"),
                        error=task.get("error"),
                    )
                )
            return files

    async def create_text_file(self, request: FileWriteRequest) -> FileMutationResponse:
        async with self._lock:
            target = self._resolve_input_file(request.filename)
            existed_before = target.exists()
            previous_bytes = target.read_bytes() if existed_before else None
            previous_overrides = self._load_metadata_overrides()

            if existed_before and not request.overwrite:
                msg = f"File already exists: {request.filename}"
                raise ValueError(msg)

            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(prepare_markdown_for_indexing(request.content), encoding="utf-8")
            self._upsert_file_metadata(
                filename=target.name,
                tenant_id=request.tenant_id,
                conversation_id=request.conversation_id,
                department_id=request.department_id,
                visibility=request.visibility,
            )
            action: Literal["created", "replaced"] = (
                "replaced" if existed_before else "created"
            )

            try:
                stats = await self._full_rebuild_unlocked(sync_mysql=True)
            except Exception:
                self._restore_input_file(target, previous_bytes)
                self._save_metadata_overrides(previous_overrides)
                raise

            return FileMutationResponse(
                filename=target.name,
                action=action,
                rebuild_mode="full",
                stats=stats,
                tenant_id=request.tenant_id,
                conversation_id=request.conversation_id,
                file_id=target.name,
                status="SUCCESS",
            )

    async def enqueue_upload_file(
        self,
        upload: UploadFile,
        overwrite: bool = False,
        tenant_id: str = "default",
        conversation_id: str | None = None,
        department_id: str = "default",
        visibility: str = "department",
    ) -> FileMutationResponse:
        if not upload.filename:
            raise ValueError("Uploaded file is missing a filename.")
        payload = await upload.read()
        normalized_name = self._target_name_for_upload(upload.filename, payload)
        target = self._resolve_input_file(normalized_name)
        async with self._lock:
            if target.exists() and not overwrite:
                msg = f"File already exists: {target.name}"
                raise ValueError(msg)
            task_id = self._create_processing_task(
                filename=target.name,
                source_filename=upload.filename,
                tenant_id=tenant_id,
                conversation_id=conversation_id,
                department_id=department_id,
                visibility=visibility,
            )
        asyncio.create_task(
            self._process_upload_task(
                task_id=task_id,
                source_filename=upload.filename,
                payload=payload,
                overwrite=overwrite,
                tenant_id=tenant_id,
                conversation_id=conversation_id,
                department_id=department_id,
                visibility=visibility,
            )
        )
        return FileMutationResponse(
            filename=target.name,
            action="replaced" if target.exists() else "created",
            rebuild_mode="full",
            stats=self._stats(),
            source_filename=upload.filename,
            tenant_id=tenant_id,
            conversation_id=conversation_id,
            file_id=target.name,
            status="PROCESSING",
            task_id=task_id,
        )

    async def _process_upload_task(
        self,
        task_id: str,
        source_filename: str,
        payload: bytes,
        overwrite: bool,
        tenant_id: str,
        conversation_id: str | None,
        department_id: str,
        visibility: str,
    ) -> None:
        try:
            result = await self.upload_file(
                _MemoryUploadFile(source_filename, payload),
                overwrite=overwrite,
                tenant_id=tenant_id,
                conversation_id=conversation_id,
                department_id=department_id,
                visibility=visibility,
            )
        except Exception as exc:
            async with self._lock:
                self._update_processing_task(task_id, status="FAILED", error=str(exc))
            return

        async with self._lock:
            self._update_processing_task(
                task_id,
                status="SUCCESS",
                filename=result.filename,
                converted=result.converted,
                conversion_type=result.conversion_type,
                warnings=result.warnings,
            )

    async def upload_file(
        self,
        upload: UploadFile,
        overwrite: bool = False,
        tenant_id: str = "default",
        conversation_id: str | None = None,
        department_id: str = "default",
        visibility: str = "department",
    ) -> FileMutationResponse:
        async with self._lock:
            if not upload.filename:
                raise ValueError("Uploaded file is missing a filename.")

            payload = await upload.read()
            normalized = self._normalize_upload(upload.filename, payload)
            target = self._resolve_input_file(normalized.filename)
            existed_before = target.exists()
            previous_bytes = target.read_bytes() if existed_before else None
            previous_overrides = self._load_metadata_overrides()

            if existed_before and not overwrite:
                msg = f"File already exists: {target.name}"
                raise ValueError(msg)

            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(
                prepare_markdown_for_indexing(normalized.content),
                encoding="utf-8",
            )
            self._upsert_file_metadata(
                filename=target.name,
                tenant_id=tenant_id,
                conversation_id=conversation_id,
                department_id=department_id,
                visibility=visibility,
            )
            action: Literal["created", "replaced"] = (
                "replaced" if existed_before else "created"
            )

            try:
                stats = await self._full_rebuild_unlocked(sync_mysql=True)
            except Exception:
                self._restore_input_file(target, previous_bytes)
                self._save_metadata_overrides(previous_overrides)
                raise

            return FileMutationResponse(
                filename=target.name,
                action=action,
                rebuild_mode="full",
                stats=stats,
                source_filename=upload.filename,
                converted=normalized.converted,
                conversion_type=normalized.conversion_type,
                warnings=normalized.warnings,
                tenant_id=tenant_id,
                conversation_id=conversation_id,
                file_id=target.name,
            )

    async def preview_upload(self, upload: UploadFile) -> FilePreviewResponse:
        async with self._lock:
            if not upload.filename:
                raise ValueError("Uploaded file is missing a filename.")
            payload = await upload.read()
            normalized = self._normalize_upload(upload.filename, payload)
            return FilePreviewResponse(
                filename=normalized.filename,
                source_filename=upload.filename,
                converted=normalized.converted,
                conversion_type=normalized.conversion_type,
                warnings=normalized.warnings,
                content=normalized.content,
            )

    async def delete_file(self, filename: str) -> FileMutationResponse:
        async with self._lock:
            target = self._resolve_input_file(filename)
            if not target.exists():
                msg = f"File not found: {filename}"
                raise FileNotFoundError(msg)

            previous_bytes = target.read_bytes()
            previous_overrides = self._load_metadata_overrides()
            file_metadata = self._file_metadata_for_response(target.name, previous_overrides)
            graph_cleanup_stats = self.neo4j_cleaner.cleanup_file(
                tenant_id=file_metadata["tenant_id"],
                file_id=file_metadata["file_id"] or target.name,
            )
            target.unlink()
            self._remove_file_metadata(target.name)
            self._remove_processing_tasks_for_file(target.name)

            try:
                stats = await self._full_rebuild_unlocked(sync_mysql=True)
            except Exception:
                target.write_bytes(previous_bytes)
                self._save_metadata_overrides(previous_overrides)
                raise
            stats["neo4j_cleanup"] = graph_cleanup_stats

            return FileMutationResponse(
                filename=target.name,
                action="deleted",
                rebuild_mode="full",
                stats=stats,
                file_id=target.name,
            )

    async def rebuild(self, verbose: bool = False, sync_mysql: bool = True) -> dict[str, Any]:
        async with self._lock:
            return await self._full_rebuild_unlocked(verbose=verbose, sync_mysql=sync_mysql)

    async def processing_task(self, task_id: str) -> dict[str, Any]:
        async with self._lock:
            task = self._load_processing_tasks().get(task_id)
            if task is None:
                raise KeyError(task_id)
            return task

    async def _reload_unlocked(self) -> dict[str, Any]:
        load_dotenv(self.root / ".env", override=False)

        config = load_config(root_dir=self.root)
        storage = create_storage(config.output_storage)
        table_provider = create_table_provider(config.table_provider, storage=storage)
        reader = DataReader(table_provider)

        documents = await reader.documents()
        entities = await reader.entities()
        communities = await reader.communities()
        community_reports = await reader.community_reports()
        text_units = await reader.text_units()
        relationships = await reader.relationships()

        try:
            covariates = await reader.covariates()
        except Exception:
            covariates = None

        self.config = config
        self.tables = LoadedTables(
            documents=documents,
            entities=entities,
            communities=communities,
            community_reports=community_reports,
            text_units=text_units,
            relationships=relationships,
            covariates=covariates,
        )
        return self._stats()

    async def _full_rebuild_unlocked(
        self,
        verbose: bool = False,
        sync_mysql: bool = True,
    ) -> dict[str, Any]:
        self.input_dir.mkdir(parents=True, exist_ok=True)
        load_dotenv(self.root / ".env", override=False)

        config = load_config(root_dir=self.root)
        try:
            validate_config_names(config)
        except SystemExit as exc:
            raise RuntimeError("GraphRAG config validation failed.") from exc

        backup_exists = False

        self._release_loaded_index()

        if self.output_backup_dir.exists():
            self._remove_tree_with_retries(self.output_backup_dir)

        if self.output_dir.exists():
            self._rename_with_retries(self.output_dir, self.output_backup_dir)
            backup_exists = True

        try:
            outputs = await build_index(
                config=config,
                method=IndexingMethod.Standard,
                is_update_run=False,
                verbose=verbose,
            )
            encountered_errors = [output for output in outputs if output.error is not None]
            if encountered_errors:
                messages = [
                    f"{output.workflow}: {output.error}" for output in encountered_errors
                ]
                raise RuntimeError("; ".join(messages))
        except Exception:
            if self.output_dir.exists():
                self._remove_tree_with_retries(self.output_dir)
            if backup_exists and self.output_backup_dir.exists():
                self._rename_with_retries(self.output_backup_dir, self.output_dir)
            await self._reload_unlocked()
            raise
        else:
            if backup_exists and self.output_backup_dir.exists():
                self._remove_tree_with_retries(self.output_backup_dir)

        stats = await self._reload_unlocked()
        if sync_mysql and self.permission_store.enabled:
            stats["mysql_sync"] = sync_graphrag_metadata_to_mysql(root=self.root)
        return stats

    def _release_loaded_index(self) -> None:
        self.config = None
        self.tables = None
        gc.collect()

    @staticmethod
    def _rename_with_retries(source: Path, target: Path, attempts: int = 8) -> None:
        last_error: OSError | None = None
        for attempt in range(attempts):
            try:
                source.rename(target)
                return
            except OSError as exc:
                last_error = exc
                gc.collect()
                time.sleep(min(0.25 * (attempt + 1), 2.0))
        raise RuntimeError(
            f"Cannot move {source} to {target}. Close any process using this project "
            "or its output/lancedb files, then rebuild again."
        ) from last_error

    @staticmethod
    def _remove_tree_with_retries(path: Path, attempts: int = 8) -> None:
        last_error: OSError | None = None
        for attempt in range(attempts):
            try:
                shutil.rmtree(path)
                return
            except FileNotFoundError:
                return
            except OSError as exc:
                last_error = exc
                gc.collect()
                time.sleep(min(0.25 * (attempt + 1), 2.0))
        raise RuntimeError(
            f"Cannot remove {path}. Close any process using this project "
            "or its output/lancedb files, then rebuild again."
        ) from last_error

    def _resolve_input_file(self, filename: str) -> Path:
        candidate = Path(filename)
        if (
            candidate.name != filename
            or candidate.is_absolute()
            or any(part in {"..", ""} for part in candidate.parts)
        ):
            msg = "Only plain filenames under input/ are allowed."
            raise ValueError(msg)
        return self.input_dir / candidate.name

    def _load_metadata_overrides(self) -> dict[str, Any]:
        if not self.metadata_overrides_path.exists():
            return {"documents": {}, "text_units": {}}
        try:
            payload = json.loads(self.metadata_overrides_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid metadata_overrides.json: {exc}") from exc
        payload.setdefault("documents", {})
        payload.setdefault("text_units", {})
        return payload

    def _save_metadata_overrides(self, payload: dict[str, Any]) -> None:
        payload.setdefault("documents", {})
        payload.setdefault("text_units", {})
        self.metadata_overrides_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def _upsert_file_metadata(
        self,
        filename: str,
        tenant_id: str = "default",
        conversation_id: str | None = None,
        department_id: str = "default",
        visibility: str = "department",
    ) -> None:
        overrides = self._load_metadata_overrides()
        document = {
            **overrides.get("documents", {}).get(filename, {}),
            "tenant_id": tenant_id or "default",
            "file_id": filename,
            "department_id": department_id or "default",
            "visibility": visibility or "department",
        }
        if conversation_id:
            document["conversation_id"] = conversation_id
            tags = list(document.get("tags") or [])
            conversation_tag = f"conversation:{conversation_id}"
            if conversation_tag not in tags:
                tags.append(conversation_tag)
            document["tags"] = tags
        else:
            document.pop("conversation_id", None)
        overrides["documents"][filename] = document
        self._save_metadata_overrides(overrides)

    def _remove_file_metadata(self, filename: str) -> None:
        overrides = self._load_metadata_overrides()
        overrides.get("documents", {}).pop(filename, None)
        self._save_metadata_overrides(overrides)

    def _document_scope_map(self) -> dict[str, dict[str, Any]]:
        assert self.tables is not None
        overrides = self._load_metadata_overrides()
        document_overrides = overrides.get("documents", {})
        scope: dict[str, dict[str, Any]] = {}
        for row in self.tables.documents.to_dict(orient="records"):
            document_id = str(row["id"])
            title = str(row.get("title", ""))
            metadata = document_overrides.get(document_id) or document_overrides.get(title)
            scope[document_id] = {
                "tenant_id": "default",
                "file_id": title or document_id,
                "department_id": "default",
                "visibility": "department",
                **(metadata or {}),
            }
        return scope

    def _file_metadata_for_response(
        self,
        filename: str,
        overrides: dict[str, Any],
    ) -> dict[str, Any]:
        metadata = overrides.get("documents", {}).get(filename, {})
        return {
            "tenant_id": str(metadata.get("tenant_id", "default")),
            "conversation_id": metadata.get("conversation_id"),
            "department_id": str(metadata.get("department_id", "default")),
            "visibility": str(metadata.get("visibility", "department")),
            "file_id": str(metadata.get("file_id", filename)),
        }

    def _load_processing_tasks(self) -> dict[str, Any]:
        if not self.processing_tasks_path.exists():
            return {}
        try:
            return json.loads(self.processing_tasks_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid processing_tasks.json: {exc}") from exc

    def _save_processing_tasks(self, payload: dict[str, Any]) -> None:
        self.processing_tasks_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def _create_processing_task(
        self,
        filename: str,
        source_filename: str,
        tenant_id: str,
        conversation_id: str | None,
        department_id: str,
        visibility: str,
    ) -> str:
        tasks = self._load_processing_tasks()
        task_id = uuid.uuid4().hex
        now = time.time()
        tasks[task_id] = {
            "task_id": task_id,
            "filename": filename,
            "source_filename": source_filename,
            "tenant_id": tenant_id,
            "conversation_id": conversation_id,
            "department_id": department_id,
            "visibility": visibility,
            "status": "PROCESSING",
            "created_at": now,
            "updated_at": now,
        }
        self._save_processing_tasks(tasks)
        return task_id

    def _update_processing_task(
        self,
        task_id: str,
        status: ProcessingStatus,
        **updates: Any,
    ) -> None:
        tasks = self._load_processing_tasks()
        task = tasks.get(task_id, {"task_id": task_id})
        task.update(updates)
        task["status"] = status
        task["updated_at"] = time.time()
        tasks[task_id] = task
        self._save_processing_tasks(tasks)

    def _remove_processing_tasks_for_file(self, filename: str) -> None:
        tasks = self._load_processing_tasks()
        kept = {
            task_id: task
            for task_id, task in tasks.items()
            if str(task.get("filename", "")) != filename
        }
        if len(kept) != len(tasks):
            self._save_processing_tasks(kept)

    @staticmethod
    def _latest_file_task_status(filename: str, tasks: dict[str, Any]) -> ProcessingStatus:
        matching = [
            task for task in tasks.values() if str(task.get("filename", "")) == filename
        ]
        if not matching:
            return "SUCCESS"
        latest = max(matching, key=lambda task: float(task.get("updated_at", 0)))
        return latest.get("status", "SUCCESS")

    @staticmethod
    def _latest_file_task_error(filename: str, tasks: dict[str, Any]) -> str | None:
        matching = [
            task for task in tasks.values() if str(task.get("filename", "")) == filename
        ]
        if not matching:
            return None
        latest = max(matching, key=lambda task: float(task.get("updated_at", 0)))
        return latest.get("error")

    def _target_name_for_upload(self, filename: str, payload: bytes) -> str:
        original_name = Path(filename).name
        stem = Path(original_name).stem or "uploaded"
        detected = _detect_upload_type(original_name, payload)
        if detected.kind in {"text", "markdown"} and detected.extension in {".txt", ".md"}:
            return original_name
        if detected.kind in {"markdown", "csv", "pdf", "docx", "xlsx"}:
            return f"{stem}.md"
        return f"{stem}.txt"

    def _normalize_upload(self, filename: str, payload: bytes) -> NormalizedUpload:
        original_name = Path(filename).name
        stem = Path(original_name).stem or "uploaded"
        warnings: list[str] = []
        detected = _detect_upload_type(original_name, payload)
        suffix = Path(original_name).suffix.lower()
        if detected.detected_by == "content" and suffix != detected.extension:
            warnings.append(
                f"File content was detected as {detected.kind}; original extension was {suffix or '(none)'}."
            )

        if detected.kind in {"text", "markdown"}:
            content = _decode_text_payload(payload)
            target_name = original_name if detected.extension in {".txt", ".md"} else f"{stem}.md"
            return NormalizedUpload(
                filename=target_name,
                content=content,
                converted=detected.kind == "markdown" and detected.extension != suffix,
                conversion_type="text",
                warnings=warnings,
            )

        if detected.kind == "pdf":
            markdown, pdf_warnings = _pdf_to_markdown(payload, original_name)
            warnings.extend(pdf_warnings)
            return NormalizedUpload(
                filename=f"{stem}.md",
                content=markdown,
                converted=True,
                conversion_type="pdf",
                warnings=warnings,
            )

        if detected.kind == "docx":
            markdown = _docx_to_markdown(payload, original_name)
            return NormalizedUpload(
                filename=f"{stem}.md",
                content=markdown,
                converted=True,
                conversion_type="docx",
                warnings=warnings,
            )

        if detected.kind == "xlsx":
            markdown = _excel_to_markdown(payload, original_name)
            return NormalizedUpload(
                filename=f"{stem}.md",
                content=markdown,
                converted=True,
                conversion_type="spreadsheet",
                warnings=warnings,
            )

        if detected.kind == "csv":
            markdown = _csv_to_markdown(payload, original_name)
            return NormalizedUpload(
                filename=f"{stem}.md",
                content=markdown,
                converted=True,
                conversion_type="csv",
                warnings=warnings,
            )

        supported = ".txt, .md, .pdf, .docx, .xlsx, .xls, .csv"
        raise ValueError(f"Unsupported file type: {suffix or '(none)'}. Supported: {supported}.")

    def _restore_input_file(
        self, target: Path, previous_bytes: bytes | None
    ) -> None:
        if previous_bytes is None:
            if target.exists():
                target.unlink()
            return
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(previous_bytes)

    def _stats(self) -> dict[str, Any]:
        file_count = (
            len([path for path in self.input_dir.iterdir() if path.is_file()])
            if self.input_dir.exists()
            else 0
        )

        if self.tables is None:
            return {
                "loaded": False,
                "root": str(self.root),
                "input_dir": str(self.input_dir),
                "output_dir": str(self.output_dir),
                "file_count": file_count,
            }

        return {
            "loaded": True,
            "root": str(self.root),
            "input_dir": str(self.input_dir),
            "output_dir": str(self.output_dir),
            "file_count": file_count,
            "counts": {
                "documents": len(self.tables.documents),
                "entities": len(self.tables.entities),
                "communities": len(self.tables.communities),
                "community_reports": len(self.tables.community_reports),
                "text_units": len(self.tables.text_units),
                "relationships": len(self.tables.relationships),
                "covariates": 0
                if self.tables.covariates is None
                else len(self.tables.covariates),
            },
        }


def to_jsonable(value: Any) -> Any:
    if isinstance(value, pd.DataFrame):
        frame = value.astype(object).where(pd.notnull(value), None)
        return frame.to_dict(orient="records")
    if isinstance(value, pd.Series):
        return [to_jsonable(item) for item in value.tolist()]
    if isinstance(value, dict):
        return {str(key): to_jsonable(item) for key, item in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [to_jsonable(item) for item in value]
    if hasattr(value, "item") and callable(value.item):
        try:
            return value.item()
        except Exception:
            pass
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    return value


def _tokenize_for_keyword_search(text: str) -> list[str]:
    """Tokenize mixed Chinese/English text for lightweight BM25 scoring."""
    normalized = text.lower()
    tokens: list[str] = []
    for match in re.finditer("[\u4e00-\u9fff]+|[a-z0-9_]+", normalized):
        piece = match.group(0)
        if re.fullmatch("[\u4e00-\u9fff]+", piece):
            tokens.extend(piece)
            if len(piece) > 1:
                tokens.extend(piece[index : index + 2] for index in range(len(piece) - 1))
        else:
            tokens.append(piece)
    return tokens


def _coerce_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple):
        return list(value)
    if hasattr(value, "tolist"):
        converted = value.tolist()
        return converted if isinstance(converted, list) else [converted]
    return [value]


def _decode_text_payload(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _detect_upload_type(filename: str, payload: bytes) -> DetectedFileType:
    suffix = Path(filename).suffix.lower()
    extension_map: dict[str, DetectedFileType] = {
        ".txt": DetectedFileType("text", ".txt", "extension"),
        ".md": DetectedFileType("markdown", ".md", "extension"),
        ".markdown": DetectedFileType("markdown", ".md", "extension"),
        ".csv": DetectedFileType("csv", ".csv", "extension"),
        ".pdf": DetectedFileType("pdf", ".pdf", "extension"),
        ".docx": DetectedFileType("docx", ".docx", "extension"),
        ".xlsx": DetectedFileType("xlsx", ".xlsx", "extension"),
        ".xls": DetectedFileType("xlsx", ".xls", "extension"),
    }

    content_detected = _detect_upload_type_from_content(payload)
    if content_detected is not None:
        if suffix not in extension_map or extension_map[suffix].kind != content_detected.kind:
            return content_detected
        return extension_map[suffix]

    if suffix in extension_map:
        return extension_map[suffix]

    supported = ".txt, .md, .pdf, .docx, .xlsx, .xls, .csv"
    raise ValueError(f"Unsupported file type: {suffix or '(none)'}. Supported: {supported}.")


def _detect_upload_type_from_content(payload: bytes) -> DetectedFileType | None:
    if payload.startswith(b"%PDF-"):
        return DetectedFileType("pdf", ".pdf", "content")

    if zipfile.is_zipfile(io.BytesIO(payload)):
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            names = set(archive.namelist())
        if "word/document.xml" in names:
            return DetectedFileType("docx", ".docx", "content")
        if "xl/workbook.xml" in names:
            return DetectedFileType("xlsx", ".xlsx", "content")

    text = _decode_text_payload(payload[:8192])
    if _looks_binary(text):
        return None
    if _looks_like_markdown(text):
        return DetectedFileType("markdown", ".md", "content")
    if _looks_like_csv(text):
        return DetectedFileType("csv", ".csv", "content")
    if text.strip():
        return DetectedFileType("text", ".txt", "content")
    return None


def _looks_binary(text: str) -> bool:
    return "\x00" in text


def _looks_like_markdown(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return False
    markdown_markers = sum(
        1
        for line in lines[:20]
        if line.startswith(("# ", "## ", "### ", "- ", "* ", "> ", "```"))
        or re.match(r"^\d+\.\s+", line)
        or ("[" in line and "](" in line)
    )
    return markdown_markers > 0


def _looks_like_csv(text: str) -> bool:
    sample = "\n".join(line for line in text.splitlines()[:8] if line.strip())
    if not sample:
        return False
    try:
        import csv

        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        rows = list(csv.reader(io.StringIO(sample), dialect))
    except Exception:
        return False
    multi_column_rows = [row for row in rows if len(row) > 1]
    return len(multi_column_rows) >= 2


def _pdf_to_markdown(payload: bytes, filename: str) -> tuple[str, list[str]]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber is required to convert PDF uploads.") from exc

    pages: list[str] = [f"# {filename}"]
    warnings: list[str] = []
    with pdfplumber.open(io.BytesIO(payload)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            tables = page.extract_tables() or []
            sections: list[str] = []
            if text:
                sections.append(text)
            for table in tables:
                rendered = _table_to_markdown(table)
                if rendered:
                    sections.append(rendered)
            if sections:
                pages.append(f"\n## Page {index}\n\n" + "\n\n".join(sections))

    if len(pages) == 1:
        ocr_text = _ocr_pdf_to_markdown(payload, filename)
        if ocr_text:
            warnings.append("PDF text extraction was empty; OCR fallback was used.")
            return ocr_text, warnings
        ocr_config = _get_ocr_config()
        msg = (
            "No extractable text was found in this PDF. It may be scanned; "
            f"{ocr_config['message']}"
        )
        raise ValueError(msg)

    if any("## Page" not in page for page in pages[1:]):
        warnings.append("Some PDF pages may have sparse text extraction results.")
    return "\n".join(pages).strip() + "\n", warnings


def _get_ocr_config() -> dict[str, Any]:
    load_dotenv(ROOT_DIR / ".env", override=False)
    provider = os.getenv("OCR_PROVIDER", os.getenv("GRAPHRAG_OCR_PROVIDER", "off")).strip().lower()
    model = os.getenv("DEEPSEEK_OCR_MODEL", "deepseek-ai/DeepSeek-OCR").strip()
    ocr_api_key = os.getenv("DEEPSEEK_OCR_API_KEY", "").strip()
    fallback_api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    api_key = ocr_api_key or fallback_api_key
    api_key_source = "DEEPSEEK_OCR_API_KEY" if ocr_api_key else "DEEPSEEK_API_KEY" if fallback_api_key else ""
    command_template = os.getenv("GRAPHRAG_OCR_COMMAND", "").strip()
    script_path = ROOT_DIR / "scripts" / "deepseek_ocr.py"
    timeout_seconds = int(os.getenv("GRAPHRAG_OCR_TIMEOUT_SECONDS", "180"))
    enabled = provider not in {"", "off", "none", "disabled"} or bool(command_template)

    if command_template:
        message = "OCR is configured through GRAPHRAG_OCR_COMMAND."
    elif provider == "deepseek":
        if script_path.exists():
            message = "OCR_PROVIDER=deepseek is configured; scripts/deepseek_ocr.py will be called."
        else:
            message = "OCR_PROVIDER=deepseek is set, but scripts/deepseek_ocr.py is missing."
    else:
        message = "OCR is not enabled. Set OCR_PROVIDER=deepseek or GRAPHRAG_OCR_COMMAND in .env."

    return {
        "provider": provider or "off",
        "enabled": enabled,
        "model": model,
        "api_key_configured": bool(api_key),
        "api_key_source": api_key_source,
        "timeout_seconds": timeout_seconds,
        "command_template": command_template,
        "command_configured": bool(command_template),
        "script_path": script_path,
        "script_exists": script_path.exists(),
        "message": message,
    }


def _build_ocr_command(input_path: Path) -> str:
    config = _get_ocr_config()
    command_template = str(config["command_template"])
    if command_template:
        return command_template.format(
            input=str(input_path),
            model=config["model"],
            api_key=os.getenv("DEEPSEEK_OCR_API_KEY", os.getenv("DEEPSEEK_API_KEY", "")),
        )

    if config["provider"] == "deepseek":
        script_path = Path(config["script_path"])
        if not script_path.exists():
            raise RuntimeError("OCR_PROVIDER=deepseek is set, but scripts/deepseek_ocr.py is missing.")
        return f'"{sys.executable}" "{script_path}" "{input_path}"'

    return ""


def _ocr_pdf_to_markdown(payload: bytes, filename: str) -> str:
    config = _get_ocr_config()
    if not config["enabled"]:
        return ""

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / filename
        input_path.write_bytes(payload)
        command = _build_ocr_command(input_path)
        if not command:
            return ""
        try:
            completed = subprocess.run(
                command,
                shell=True,
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=int(config["timeout_seconds"]),
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("OCR command timed out.") from exc
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            detail = f": {stderr}" if stderr else ""
            raise RuntimeError(f"OCR command failed{detail}") from exc

    text = completed.stdout.strip()
    if not text:
        return ""
    if text.lstrip().startswith("#"):
        return text + "\n"
    return f"# {filename}\n\n## OCR\n\n{text}\n"


def _docx_to_markdown(payload: bytes, filename: str) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required to convert DOCX uploads.") from exc

    document = docx.Document(io.BytesIO(payload))
    lines = [f"# {filename}"]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if "heading 1" in style_name:
            lines.append(f"\n# {text}")
        elif "heading 2" in style_name:
            lines.append(f"\n## {text}")
        elif "heading 3" in style_name:
            lines.append(f"\n### {text}")
        else:
            lines.append(text)

    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rendered = _table_to_markdown(rows)
        if rendered:
            lines.append(f"\n## Table {index}\n\n{rendered}")

    content = "\n\n".join(lines).strip()
    if content == f"# {filename}":
        raise ValueError("No readable text was found in the DOCX file.")
    return content + "\n"


def _excel_to_markdown(payload: bytes, filename: str) -> str:
    sheets = pd.read_excel(io.BytesIO(payload), sheet_name=None, dtype=object)
    sections = [f"# {filename}"]
    for sheet_name, frame in sheets.items():
        section = _dataframe_to_statements(frame, f"Sheet: {sheet_name}")
        if section:
            sections.append(section)
    if len(sections) == 1:
        raise ValueError("No readable rows were found in the spreadsheet.")
    return "\n\n".join(sections).strip() + "\n"


def _csv_to_markdown(payload: bytes, filename: str) -> str:
    text = _decode_text_payload(payload)
    frame = pd.read_csv(io.StringIO(text), dtype=object)
    section = _dataframe_to_statements(frame, filename)
    if not section:
        raise ValueError("No readable rows were found in the CSV file.")
    return f"# {filename}\n\n{section}\n"


def _sse_event(event: str, payload: Any) -> str:
    data = json.dumps(to_jsonable(payload), ensure_ascii=False)
    return f"event: {event}\ndata: {data}\n\n"


def stringify_stream_answer(answer: Any) -> str:
    if answer is None:
        return ""
    if isinstance(answer, str):
        return clean_customer_answer(answer)
    return clean_customer_answer(json.dumps(to_jsonable(answer), ensure_ascii=False))


def clean_customer_answer(answer: Any) -> Any:
    if not isinstance(answer, str):
        return answer
    cleaned = re.sub(r"\s*\[Data:[^\]]+\]", "", answer)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def split_stream_text(text: str, chunk_size: int = 48) -> list[str]:
    if not text:
        return [""]
    return [text[index : index + chunk_size] for index in range(0, len(text), chunk_size)]


def prepare_markdown_for_indexing(content: str, max_section_chars: int = 5000) -> str:
    sections = split_markdown_sections(content, max_section_chars=max_section_chars)
    if not sections:
        return content
    rendered: list[str] = []
    for index, section in enumerate(sections, start=1):
        path = " > ".join(section["path"]) if section["path"] else "Document"
        rendered.append(f"<!-- chunk:{index} path:{path} -->\n{section['text'].strip()}")
    return "\n\n".join(rendered).strip() + "\n"


def split_markdown_sections(content: str, max_section_chars: int = 5000) -> list[dict[str, Any]]:
    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    sections: list[dict[str, Any]] = []
    current_lines: list[str] = []
    heading_stack: dict[int, str] = {}

    def current_path() -> list[str]:
        return [heading_stack[level] for level in sorted(heading_stack)]

    def flush() -> None:
        nonlocal current_lines
        text = "\n".join(current_lines).strip()
        if not text:
            current_lines.clear()
            return
        path = current_path()
        for part in _split_large_markdown_section(text, max_section_chars):
            sections.append({"path": path, "text": part})
        current_lines.clear()

    for line in lines:
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush()
            level = len(match.group(1))
            heading_stack = {
                existing_level: title
                for existing_level, title in heading_stack.items()
                if existing_level < level
            }
            heading_stack[level] = match.group(2).strip()
        current_lines.append(line)
    flush()
    return sections


def _split_large_markdown_section(text: str, max_section_chars: int) -> list[str]:
    if len(text) <= max_section_chars:
        return [text]

    blocks = re.split(r"\n{2,}", text)
    chunks: list[str] = []
    current: list[str] = []
    current_size = 0
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        extra = len(block) + 2
        if current and current_size + extra > max_section_chars:
            chunks.append("\n\n".join(current))
            current = []
            current_size = 0
        if len(block) > max_section_chars:
            chunks.extend(
                block[index : index + max_section_chars]
                for index in range(0, len(block), max_section_chars)
            )
            continue
        current.append(block)
        current_size += extra
    if current:
        chunks.append("\n\n".join(current))
    return chunks


def _dataframe_to_statements(frame: pd.DataFrame, title: str) -> str:
    if frame.empty:
        return ""
    cleaned = frame.dropna(how="all").copy()
    if cleaned.empty:
        return ""

    columns = _normalize_table_columns(cleaned.columns)
    cleaned.columns = columns
    lines = [f"## {title}"]
    for row_index, row in enumerate(cleaned.to_dict(orient="records"), start=1):
        facts: list[str] = []
        subject = ""
        for column in columns:
            value = row.get(column)
            if pd.isna(value):
                continue
            value_text = str(value).strip()
            if not value_text:
                continue
            if not subject:
                subject = value_text
            facts.append(f"{column}是{value_text}")
        if not facts:
            continue
        prefix = subject or f"第{row_index}行"
        lines.append(f"- {prefix}：" + "，".join(facts) + "。")
    return "\n".join(lines) if len(lines) > 1 else ""


def _normalize_table_columns(columns: Any) -> list[str]:
    normalized: list[str] = []
    seen: defaultdict[str, int] = defaultdict(int)
    for index, column in enumerate(columns):
        name = str(column).strip()
        if not name or name.lower().startswith("unnamed:"):
            name = f"列{index + 1}"
        seen[name] += 1
        if seen[name] > 1:
            name = f"{name}{seen[name]}"
        normalized.append(name)
    return normalized


def _table_to_markdown(table: list[list[Any]]) -> str:
    rows = [
        ["" if cell is None else str(cell).replace("\n", " ").strip() for cell in row]
        for row in table
        if row and any(cell is not None and str(cell).strip() for cell in row)
    ]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    if not body:
        body = [[""] * width]
    return "\n".join([
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
        *["| " + " | ".join(row) + " |" for row in body],
    ])


ROOT_DIR = Path(os.getenv("GRAPHRAG_ROOT", Path(__file__).resolve().parent))
service = GraphRAGService(ROOT_DIR)
FRONTEND_DIR = ROOT_DIR / "frontend"

app = FastAPI(
    title="GraphRAG Query API",
    version="1.1.0",
    description="HTTP wrapper for querying and maintaining the local GraphRAG index.",
    docs_url=None,
    redoc_url=None,
    openapi_url=None,
)

cors_origins = [
    item.strip()
    for item in os.getenv("GRAPHRAG_CORS_ORIGINS", "*").split(",")
    if item.strip()
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins or ["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

if FRONTEND_DIR.exists():
    app.mount("/ui", StaticFiles(directory=FRONTEND_DIR, html=True), name="ui")


@app.on_event("startup")
async def startup_event() -> None:
    await service.reload()


@app.get("/health")
async def health() -> dict[str, Any]:
    return service._stats()


@app.get("/")
async def index() -> FileResponse:
    index_path = FRONTEND_DIR / "index.html"
    if not index_path.exists():
        raise HTTPException(status_code=404, detail="Frontend is not built.")
    return FileResponse(index_path)


@app.get("/methods")
async def methods() -> dict[str, list[str]]:
    return {"methods": ["basic", "local", "global", "drift", "hybrid"]}


@app.get("/api/v1/admin/ocr-config", response_model=OcrConfigResponse)
async def ocr_config() -> OcrConfigResponse:
    config = _get_ocr_config()
    return OcrConfigResponse(
        provider=config["provider"],
        enabled=config["enabled"],
        model=config["model"],
        api_key_configured=config["api_key_configured"],
        api_key_source=config["api_key_source"],
        timeout_seconds=config["timeout_seconds"],
        command_configured=config["command_configured"],
        script_exists=config["script_exists"],
        message=config["message"],
    )


@app.get("/files", response_model=list[FileInfo])
@app.get("/api/v1/admin/files", response_model=list[FileInfo])
async def list_files() -> list[FileInfo]:
    try:
        return await service.list_files()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"List files failed: {exc}") from exc


@app.post("/files/text", response_model=FileMutationResponse)
@app.post("/api/v1/admin/files/text", response_model=FileMutationResponse)
async def create_text_file(request: FileWriteRequest) -> FileMutationResponse:
    try:
        return await service.create_text_file(request)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Create file failed: {exc}") from exc


@app.post("/files/upload", response_model=FileMutationResponse)
@app.post("/api/v1/admin/files/upload", response_model=FileMutationResponse)
async def upload_file(
    file: UploadFile = File(...),
    overwrite: bool = Form(False),
    tenant_id: str = Form("default"),
    conversation_id: str | None = Form(None),
    department_id: str = Form("default"),
    visibility: str = Form("department"),
) -> FileMutationResponse:
    try:
        return await service.enqueue_upload_file(
            file,
            overwrite=overwrite,
            tenant_id=tenant_id,
            conversation_id=conversation_id,
            department_id=department_id,
            visibility=visibility,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Upload failed: {exc}") from exc


@app.get("/files/tasks/{task_id}")
@app.get("/api/v1/admin/files/tasks/{task_id}")
async def file_task(task_id: str) -> dict[str, Any]:
    try:
        return await service.processing_task(task_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"Task not found: {task_id}") from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Task lookup failed: {exc}") from exc


@app.post("/files/preview", response_model=FilePreviewResponse)
@app.post("/api/v1/admin/files/preview", response_model=FilePreviewResponse)
async def preview_file(file: UploadFile = File(...)) -> FilePreviewResponse:
    try:
        return await service.preview_upload(file)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Preview failed: {exc}") from exc


@app.delete("/files/{filename}", response_model=FileMutationResponse)
@app.delete("/api/v1/admin/files/{filename}", response_model=FileMutationResponse)
async def delete_file(filename: str) -> FileMutationResponse:
    try:
        return await service.delete_file(filename)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Delete failed: {exc}") from exc


@app.post("/rebuild")
async def rebuild_index(request: RebuildRequest | None = None) -> dict[str, Any]:
    request = request or RebuildRequest()
    try:
        return await service.rebuild(
            verbose=request.verbose,
            sync_mysql=request.sync_mysql,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Rebuild failed: {exc}") from exc


@app.post("/reload")
async def reload_index() -> dict[str, Any]:
    try:
        return await service.reload()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Reload failed: {exc}") from exc


@app.post("/query", response_model=QueryResponse)
async def query(request: QueryRequest) -> QueryResponse:
    try:
        return await service.query(request)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Query failed: {exc}") from exc


@app.post("/api/v1/chat/stream")
async def chat_stream(request: ChatStreamRequest) -> StreamingResponse:
    async def event_generator():
        try:
            yield _sse_event("status", {"stage": "retrieving"})
            if request.method == "hybrid":
                async with service._lock:
                    if service.config is None or service.tables is None:
                        await service._reload_unlocked()
                    context_text, context_payload = service._hybrid_context_unlocked(
                        QueryRequest(
                            question=request.message,
                            method=request.method,
                            tenant_id=request.tenant_id,
                            conversation_id=request.conversation_id,
                            include_context=request.include_context,
                            response_type=request.response_type,
                            hybrid_top_k=request.hybrid_top_k,
                            hybrid_graph_k=request.hybrid_graph_k,
                        )
                    )
                    if request.include_context:
                        context_json = to_jsonable(context_payload)
                    else:
                        context_json = None
                yield _sse_event("status", {"stage": "generating"})
                if context_text:
                    async for chunk in service._stream_basic_answer(
                        query=request.message,
                        context_text=context_text,
                        response_type=request.response_type,
                    ):
                        if chunk:
                            yield _sse_event("delta", {"text": chunk})
                if request.include_context:
                    yield _sse_event("context", context_json)
                yield _sse_event("done", {})
                return

            response = await service.query(
                QueryRequest(
                    question=request.message,
                    method=request.method,
                    tenant_id=request.tenant_id,
                    conversation_id=request.conversation_id,
                    include_context=request.include_context,
                    response_type=request.response_type,
                    hybrid_top_k=request.hybrid_top_k,
                    hybrid_graph_k=request.hybrid_graph_k,
                )
            )
            yield _sse_event("status", {"stage": "generating"})
            answer = stringify_stream_answer(response.answer)
            for chunk in split_stream_text(answer):
                yield _sse_event("delta", {"text": chunk})
                await asyncio.sleep(0)
            if request.include_context:
                yield _sse_event("context", to_jsonable(response.context))
            yield _sse_event("done", {})
        except Exception as exc:
            yield _sse_event("error", {"detail": str(exc)})

    return StreamingResponse(event_generator(), media_type="text/event-stream")


if __name__ == "__main__":
    import uvicorn

    host = os.getenv("GRAPHRAG_HOST", "127.0.0.1")
    port = int(os.getenv("GRAPHRAG_PORT", "8000"))
    uvicorn.run("app:app", host=host, port=port, reload=False)
